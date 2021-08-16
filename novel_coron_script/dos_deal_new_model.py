# -*- coding: utf-8 -*-
"""
Created on Mon May 10 09:39:49 2021

@author: qic
"""
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt,Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from glob import glob
import datetime
import sys,os

#字符替换
def replace_word(doc,new_text,old_text,*check):
    list_table = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if not cell.tables:
                    cell.text = cell.text.replace(old_text, new_text)
                    if check[0] == 1:
                        if 'ORF1ab基因设计的引物探针区域未发生突变。N基因扩增正向引物5’末端有1个单碱基突变（G28881T共1482条,G28881W共1条,G28881A共1条，G28881G共6条，另外共42条序列在该位置测序未覆盖）。原理上单碱基的突变不会对检测结果造成影响。' in cell.text:
                            cell.text = cell.text.split('：')[0] + '：' + check[1]
                else:
                    list_table.append(cell.tables[0])
                    
    if check[0] == 2:
        return list_table

#突变表更改
def table_change(new_table,list_info):
    dict_info = {}
    #增加和突变位点一样多的行数
    for i in range(len(list_info)-len(new_table.rows)+1):
        new_table.add_row()
    
    #替换表格字符
    for p in range(1,len(new_table.rows)):
        for i in range(len(new_table.rows[p].cells)):
            chr_text = new_table.rows[p].cells[i].text
            chr_text = list_info[p-1][i]
            #设置表格字符格式
            paragraph = new_table.rows[p].cells[i].paragraphs[0]
            run = paragraph.add_run(chr_text)
            run.font.name = '宋体'
            run.font.size = Pt(9)
            paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if list_info[p-1][-1] == 'ORF' or list_info[p-1][-1] == 'N':
                run.font.color.rgb = RGBColor(255,0,0)
                
        if len(list_info[p-1]) > 7:
            info_variant = list_info[p-1]
            if info_variant[-1] not in dict_info:dict_info[info_variant[-1]] = []
            dict_info[info_variant[-1]].append([info_variant[1],info_variant[-3],info_variant[-2]])
    return dict_info

#更改结论
def conclusion_make(gene,total_num,dict_info,word):
    if gene in dict_info:
        mut_info = dict_info[gene]
        gene = 'ORF1ab' if gene == 'ORF' else gene
        word += '%s基因设计的引物探针区域发生%s个突变（'%(gene,len(mut_info))
        for key in mut_info:
            mut_list = key[0].split(',')
            num_list = [int(i) for i in key[1].split(',')]
            for i in range(len(mut_list)):
                word += '%s共%s条,'%(mut_list[i],num_list[i])
            word += '其中%s条没有发生突变,还有%s条未覆盖;'%(int(key[2])-sum(num_list),total_num-int(key[2]))
        word += '）。'
    else:
        gene = 'ORF1ab' if gene == 'ORF' else gene
        word += '%s基因设计的引物探针区域未发生突变。'%gene
    return word

def read_mut(file):
    list_info = []
    with open(file,'r',encoding='utf-8') as f:
        for line in f:
            if line.startswith('位置'):continue
            list_info.append(line.strip().split('\t'))
    return list_info

#添加图片
def pic_set(img,picture):
    pic = img.add_picture(picture)
    pic.height=Cm(11.34)
    pic.width=Cm(23.74)

def pic_add(doc,species,path_img):
    list_img = glob('%s/%s/*.png'%(path_img,species))
    if list_img:
        for i in list_img:
            img = doc.tables[0].cell(6,3).paragraphs[1].add_run()
            pic_set(img,i)
        return 1
    else:
        return 0

#添加序列编号
def acc_id_add(doc,file):
    with open(file,'r') as f:
        for line in f:
            line = line.strip().split('、')
    doc.tables[0].cell(7,3).paragraphs[1].text = doc.tables[0].cell(7,3).paragraphs[1].text.rstrip('。') + '、'.join(line[:20]) + '等。'

def main():
    now = datetime.datetime.now()
    #模板word
    model_fatovcf_docx = '/project/user/chenqi/test/test_nothing/novel_coron_script/bin/model_mutation_fatovcf.docx'
    model_samtools_docx = '/project/user/chenqi/test/test_nothing/novel_coron_script/bin/model_samtools.docx'
    #读取文件
    if sys.argv[1] == 'samtools':
        doc = Document(model_samtools_docx)
    elif sys.argv[1] == 'fatovcf':
        doc = Document(model_fatovcf_docx)
    else:
        print('only samtools or fatovcf can be choose!')
        sys.exit(0)

    file = sys.argv[2]
    species = file.split('/')[-1].rstrip('.fasta.txt')
    list_info = read_mut(file)
    #第三个参数应该为该珠系序列总数
    total_num = int(sys.argv[3])
    
    #添加图片
    path_img = '/project/user/chenqi/test/test_nothing/static/img'
    check_img = 0
    if os.path.exists(path_img + '/%s'%species):
        check_img = pic_add(doc,species,path_img)

    #添加序列编号
    check_acc = 0
    if os.path.exists('%s.acc.txt'%species):
        acc_id_add(doc,'%s.acc.txt'%species)
        check_acc = 1

    list_table = replace_word(doc,species,'B.1.617+',2)
    new_table = list_table[-1]
    dict_info = table_change(new_table,list_info)
    #结论获取
    word = ''
    word = conclusion_make('ORF',total_num,dict_info,word)
    word = conclusion_make('N',total_num,dict_info,word)
    #结论更改
    replace_word(doc,str(total_num),'1532',1,word)
    #日期更改
    today_date = [str(now.year),str(now.month),str(now.day)]
    replace_word(doc,'.'.join(today_date),'2021.4.28',0)
    #国家更改
    #第四个参数应该为国家，如果没有默认为原始的模板上面的国家--印度
    replace_word(doc,'','印度',0)
    # if len(sys.argv) == 5:
    #     country = sys.argv[4]
    #     replace_word(doc,country,'印度',0)
    # else:
    #     country = ''
    #     replace_word(doc,country,'印度',0)
    
    name = '圣湘'
    if len(sys.argv) == 5:
        if sys.argv[4] == 'C' or sys.argv[4] == 'S':
            name = '圣湘'
        elif sys.argv[4] == 'G':
            name = 'sansure'
    #储存新文件
    if check_img and check_acc:
        doc.save('%s_附件2 评价报告表_%s_%s.docx'%(name,species,'-'.join(today_date)))
    elif not check_img and not check_acc:
        doc.save('%s_附件2 评价报告表_%s_%s_img_acc_not.docx'%(name,species,'-'.join(today_date)))
    elif not check_img:
        doc.save('%s_附件2 评价报告表_%s_%s_img_not.docx'%(name,species,'-'.join(today_date)))
    else:
        doc.save('%s_附件2 评价报告表_%s_%s_acc_not.docx'%(name,species,'-'.join(today_date)))
        

help_info = """<usage>:
    python3 %s <samtools|fatovcf> <mutation file> <num of sequnce:int> <prob:C|G|S>"""%sys.argv[0]
if len(sys.argv) <4 or len(sys.argv)>5:
    print(help_info)
    sys.exit(0)
main()
    
